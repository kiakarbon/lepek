import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import io
import json
import sqlite3
from pathlib import Path
import hashlib
import secrets
import re
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
import matplotlib.pyplot as plt
from email_validator import validate_email, EmailNotValidError

# ============================================
# KONFIGURASI AWAL
# ============================================

# Konfigurasi halaman
st.set_page_config(
    page_title="NaNote - Catatan Praktikum & Kalkulasi PSA",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Custom untuk tampilan yang lebih baik
st.markdown("""
<style>
    .main-title {
        text-align: center;
        color: #1E3A8A;
        font-size: 2.8rem;
        margin-bottom: 1rem;
    }
    .sub-title {
        text-align: center;
        color: #4B5563;
        font-size: 1.2rem;
        margin-bottom: 2rem;
    }
    .section-header {
        color: #1D546D;
        border-left: 5px solid #3B82F6;
        padding-left: 15px;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .data-box {
        background-color: #C0C9EE;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
        border: 1px solid #E5E7EB;
    }
    .result-box {
        background-color: #C0C9EE;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        border-left: 5px solid #10B981;
    }
    .warning-box {
        background-color: #FEF3C7;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        border-left: 5px solid #F59E0B;
    }
    .success-box {
        background-color: #D1FAE5;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        border-left: 5px solid #10B981;
    }
    .error-box {
        background-color: #FEE2E2;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        border-left: 5px solid #EF4444;
    }
    .stButton button {
        background-color: #3B82F6;
        color: white;
        border-radius: 5px;
        padding: 10px 24px;
        font-weight: bold;
        border: none;
    }
    .stDownloadButton button {
        background-color: #10B981;
        color: white;
        border-radius: 5px;
        padding: 10px 24px;
        font-weight: bold;
    }
    .footer {
        text-align: center;
        margin-top: 3rem;
        color: #6B7280;
        font-size: 0.9rem;
    }
    .stats-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 10px;
        padding: 20px;
        color: white;
        margin: 10px 0;
    }
    .login-container {
        max-width: 400px;
        margin: 50px auto;
        padding: 30px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        background-color: white;
    }
    .user-profile {
        text-align: center;
        padding: 15px;
        border-bottom: 1px solid #E5E7EB;
        margin-bottom: 20px;
    }
    .user-avatar {
        width: 80px;
        height: 80px;
        border-radius: 50%;
        margin: 0 auto 15px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        display: flex;
        align-items: center;
        justify-content: center;
        color: white;
        font-size: 32px;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# FUNGSI KEAMANAN DAN HASHING
# ============================================

def generate_salt():
    """Generate salt untuk hashing password"""
    return secrets.token_hex(16)

def hash_password(password, salt=None):
    """Hash password dengan salt"""
    if salt is None:
        salt = generate_salt()
    
    # Gabungkan password dan salt
    salted_password = password + salt
    
    # Hash menggunakan SHA-256
    hashed = hashlib.sha256(salted_password.encode()).hexdigest()
    
    return hashed, salt

def verify_password(password, hashed_password, salt):
    """Verifikasi password"""
    new_hash, _ = hash_password(password, salt)
    return new_hash == hashed_password

def is_valid_email(email):
    """Validasi format email"""
    try:
        # Validasi email
        valid = validate_email(email)
        return True, valid.email
    except EmailNotValidError as e:
        return False, str(e)

def is_valid_password(password):
    """Validasi kekuatan password"""
    errors = []
    
    if len(password) < 8:
        errors.append("Password minimal 8 karakter")
    
    if not re.search(r"[A-Z]", password):
        errors.append("Password harus mengandung minimal 1 huruf besar")
    
    if not re.search(r"[a-z]", password):
        errors.append("Password harus mengandung minimal 1 huruf kecil")
    
    if not re.search(r"[0-9]", password):
        errors.append("Password harus mengandung minimal 1 angka")
    
    if not re.search(r"[!@#$%^&*()\-_=+{};:,<.>]", password):
        errors.append("Password harus mengandung minimal 1 karakter khusus")
    
    return len(errors) == 0, errors

# ============================================
# FUNGSI DATABASE
# ============================================

def init_database():
    """Inisialisasi database SQLite dengan tabel untuk autentikasi"""
    conn = sqlite3.connect('nanote.db', check_same_thread=False)
    c = conn.cursor()
    
    # Tabel untuk pengguna
    c.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            salt TEXT NOT NULL,
            full_name TEXT,
            institution TEXT,
            role TEXT DEFAULT 'user',
            is_active INTEGER DEFAULT 1,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            last_login DATETIME,
            login_count INTEGER DEFAULT 0
        )
    ''')
    
    # Tabel untuk riwayat pengguna
    c.execute('''
        CREATE TABLE IF NOT EXISTS user_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            activity_type TEXT,
            activity_data TEXT,
            ip_address TEXT,
            user_agent TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Tabel untuk catatan praktikum
    c.execute('''
        CREATE TABLE IF NOT EXISTS catatan_praktikum (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            judul TEXT,
            praktikan TEXT,
            mata_praktikum TEXT,
            tanggal TEXT,
            kelompok TEXT,
            pic TEXT,
            tujuan TEXT,
            alat_bahan TEXT,
            prosedur TEXT,
            hasil TEXT,
            analisis TEXT,
            kesimpulan TEXT,
            is_public INTEGER DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Tabel untuk hasil PSA
    c.execute('''
        CREATE TABLE IF NOT EXISTS hasil_psa (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            data_input TEXT,
            diameter_rata REAL,
            pdi_rata REAL,
            total_vol REAL,
            kualitas TEXT,
            distribusi TEXT,
            jumlah_partikel INTEGER,
            is_public INTEGER DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Tabel untuk statistik penggunaan
    c.execute('''
        CREATE TABLE IF NOT EXISTS usage_stats (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            page_views INTEGER DEFAULT 0,
            catatan_created INTEGER DEFAULT 0,
            psa_calculated INTEGER DEFAULT 0,
            last_active DATETIME,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Tabel untuk sesi login
    c.execute('''
        CREATE TABLE IF NOT EXISTS user_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            session_token TEXT UNIQUE,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            expires_at DATETIME,
            is_active INTEGER DEFAULT 1,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Buat admin default jika belum ada
    c.execute("SELECT COUNT(*) FROM users WHERE username = 'admin'")
    if c.fetchone()[0] == 0:
        # Password: Admin123!
        password_hash, salt = hash_password("Admin123!")
        c.execute('''
            INSERT INTO users (username, email, password_hash, salt, full_name, role)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', ('admin', 'admin@nanote.app', password_hash, salt, 'Administrator', 'admin'))
    
    conn.commit()
    return conn

# ============================================
# FUNGSI AUTENTIKASI
# ============================================

def register_user(conn, username, email, password, full_name="", institution=""):
    """Registrasi pengguna baru"""
    c = conn.cursor()
    
    # Validasi input
    if not username or not email or not password:
        return False, "Semua field harus diisi"
    
    # Validasi email
    email_valid, email_msg = is_valid_email(email)
    if not email_valid:
        return False, f"Email tidak valid: {email_msg}"
    
    # Validasi password
    password_valid, password_errors = is_valid_password(password)
    if not password_valid:
        return False, f"Password lemah: {', '.join(password_errors)}"
    
    # Cek apakah username sudah ada
    c.execute("SELECT id FROM users WHERE username = ?", (username,))
    if c.fetchone():
        return False, "Username sudah digunakan"
    
    # Cek apakah email sudah ada
    c.execute("SELECT id FROM users WHERE email = ?", (email,))
    if c.fetchone():
        return False, "Email sudah terdaftar"
    
    # Hash password
    password_hash, salt = hash_password(password)
    
    # Simpan pengguna ke database
    try:
        c.execute('''
            INSERT INTO users (username, email, password_hash, salt, full_name, institution)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (username, email, password_hash, salt, full_name, institution))
        
        user_id = c.lastrowid
        
        # Buat statistik awal
        c.execute('''
            INSERT INTO usage_stats (user_id)
            VALUES (?)
        ''', (user_id,))
        
        conn.commit()
        
        # Log aktivitas
        log_activity(conn, user_id, 'user_registered', {
            'username': username,
            'email': email
        })
        
        return True, f"Registrasi berhasil! Silakan login dengan username: {username}"
    
    except Exception as e:
        return False, f"Error saat registrasi: {str(e)}"

def login_user(conn, username, password):
    """Login pengguna"""
    c = conn.cursor()
    
    # Cari pengguna berdasarkan username atau email
    c.execute('''
        SELECT id, username, password_hash, salt, role, is_active 
        FROM users 
        WHERE username = ? OR email = ?
    ''', (username, username))
    
    user = c.fetchone()
    
    if not user:
        return False, "Username/email atau password salah"
    
    user_id, db_username, password_hash, salt, role, is_active = user
    
    # Cek apakah akun aktif
    if not is_active:
        return False, "Akun dinonaktifkan. Hubungi administrator."
    
    # Verifikasi password
    if not verify_password(password, password_hash, salt):
        # Update login attempt count
        log_activity(conn, user_id, 'login_failed', {
            'username': db_username,
            'reason': 'wrong_password'
        })
        return False, "Username/email atau password salah"
    
    # Update last login dan login count
    c.execute('''
        UPDATE users 
        SET last_login = CURRENT_TIMESTAMP, 
            login_count = login_count + 1 
        WHERE id = ?
    ''', (user_id,))
    
    # Buat session token
    session_token = secrets.token_urlsafe(32)
    expires_at = datetime.now().timestamp() + (24 * 60 * 60)  # 24 jam
    
    c.execute('''
        INSERT INTO user_sessions (user_id, session_token, expires_at)
        VALUES (?, ?, datetime(?, 'unixepoch'))
    ''', (user_id, session_token, expires_at))
    
    conn.commit()
    
    # Log aktivitas login berhasil
    log_activity(conn, user_id, 'login_success', {
        'username': db_username,
        'session_token': session_token[:10] + "..."
    })
    
    return True, {
        'user_id': user_id,
        'username': db_username,
        'role': role,
        'session_token': session_token
    }

def logout_user(conn, user_id, session_token):
    """Logout pengguna"""
    c = conn.cursor()
    
    # Nonaktifkan session
    c.execute('''
        UPDATE user_sessions 
        SET is_active = 0 
        WHERE user_id = ? AND session_token = ?
    ''', (user_id, session_token))
    
    conn.commit()
    
    # Log aktivitas
    log_activity(conn, user_id, 'logout', {
        'session_token': session_token[:10] + "..."
    })
    
    return True

def get_current_user(conn):
    """Mendapatkan informasi pengguna saat ini dari session"""
    if 'user' not in st.session_state:
        return None
    
    user_info = st.session_state.user
    c = conn.cursor()
    
    # Verifikasi session masih valid
    c.execute('''
        SELECT us.is_active, us.expires_at, u.username, u.role, u.full_name, u.email
        FROM user_sessions us
        JOIN users u ON us.user_id = u.id
        WHERE us.user_id = ? 
          AND us.session_token = ? 
          AND us.is_active = 1
          AND us.expires_at > CURRENT_TIMESTAMP
    ''', (user_info['user_id'], user_info['session_token']))
    
    session = c.fetchone()
    
    if not session:
        # Session tidak valid, clear session state
        if 'user' in st.session_state:
            del st.session_state.user
        return None
    
    return {
        'user_id': user_info['user_id'],
        'username': session[2],
        'role': session[3],
        'full_name': session[4],
        'email': session[5],
        'session_token': user_info['session_token']
    }

def log_activity(conn, user_id, activity_type, activity_data=None, ip_address=None, user_agent=None):
    """Mencatat aktivitas pengguna"""
    c = conn.cursor()
    
    # Dapatkan IP address dan user agent jika tersedia
    if not ip_address:
        try:
            # Untuk Streamlit Cloud, kita bisa mendapatkan IP dari request context
            import streamlit.runtime.scriptrunner.script_run_context as context
            ctx = context.get_script_run_ctx()
            if ctx:
                ip_address = ctx.request.remote_ip if hasattr(ctx.request, 'remote_ip') else 'unknown'
                user_agent = ctx.request.headers.get('User-Agent', 'unknown') if hasattr(ctx.request, 'headers') else 'unknown'
        except:
            ip_address = 'unknown'
            user_agent = 'unknown'
    
    c.execute('''
        INSERT INTO user_history (user_id, activity_type, activity_data, ip_address, user_agent)
        VALUES (?, ?, ?, ?, ?)
    ''', (user_id, activity_type, json.dumps(activity_data) if activity_data else None, ip_address, user_agent))
    
    # Update last_active di usage_stats
    c.execute('''
        UPDATE usage_stats 
        SET last_active = CURRENT_TIMESTAMP 
        WHERE user_id = ?
    ''', (user_id,))
    
    # Update page views untuk aktivitas page_view
    if activity_type == 'page_view':
        c.execute('''
            UPDATE usage_stats 
            SET page_views = page_views + 1 
            WHERE user_id = ?
        ''', (user_id,))
    
    conn.commit()

# ============================================
# FUNGSI UTAMA APLIKASI
# ============================================

def kalkulasi_hasil_psa(pdi, vol, diameter):
    """Kalkulasi hasil PSA berdasarkan parameter input"""
    try:
        pdi_array = np.array(pdi)
        vol_array = np.array(vol)
        diameter_array = np.array(diameter)
        
        total_vol = np.sum(vol_array)
        if total_vol == 0:
            return None
            
        diameter_rata = np.sum(diameter_array * vol_array) / total_vol
        pdi_rata = np.sum(pdi_array * vol_array) / total_vol
        
        distribusi = {
            '<10 nm': np.sum(vol_array[diameter_array < 10]),
            '10-50 nm': np.sum(vol_array[(diameter_array >= 10) & (diameter_array < 50)]),
            '50-100 nm': np.sum(vol_array[(diameter_array >= 50) & (diameter_array < 100)]),
            '100-500 nm': np.sum(vol_array[(diameter_array >= 100) & (diameter_array < 500)]),
            '>500 nm': np.sum(vol_array[diameter_array >= 500])
        }
        
        if pdi_rata < 0.1:
            kualitas = "Sangat Baik (Monodispers)"
            warna_kualitas = "green"
        elif pdi_rata < 0.2:
            kualitas = "Baik"
            warna_kualitas = "lightgreen"
        elif pdi_rata < 0.3:
            kualitas = "Cukup"
            warna_kualitas = "orange"
        else:
            kualitas = "Kurang (Polydispers Tinggi)"
            warna_kualitas = "red"
        
        hasil = {
            'diameter_rata': float(diameter_rata),
            'pdi_rata': float(pdi_rata),
            'total_vol': float(total_vol),
            'distribusi': distribusi,
            'kualitas': kualitas,
            'warna_kualitas': warna_kualitas,
            'jumlah_partikel': len(pdi)
        }
        
        return hasil
        
    except Exception as e:
        st.error(f"Error dalam perhitungan: {str(e)}")
        return None

def buat_file_word(catatan_data):
    doc = Document()
    doc.add_heading('Catatan Praktikum NaNote', 0)
    doc.add_paragraph(f"Judul: {catatan_data.get('judul', 'Tanpa Judul')}")
    doc.add_paragraph(f"Tanggal: {catatan_data.get('tanggal', datetime.now().strftime('%Y-%m-%d'))}")
    doc.add_paragraph(f"Praktikan: {catatan_data.get('praktikan', 'Tidak Diketahui')}")
    doc.add_paragraph(f"Mata Praktikum: {catatan_data.get('mata_praktikum', 'Tidak Diketahui')}")
    doc.add_paragraph()
    
    doc.add_heading('Isi Catatan', level=1)
    for bagian, isi in catatan_data.get('isi', {}).items():
        doc.add_heading(bagian, level=2)
        doc.add_paragraph(isi)
    
    if 'data_psa' in catatan_data:
        doc.add_heading('Data PSA', level=1)
        for key, value in catatan_data['data_psa'].items():
            doc.add_paragraph(f"{key}: {value}")
    
    return doc

def buat_file_pdf(hasil_psa, data_input):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 750, "LAPORAN HASIL PSA NANOMATERIAL")
    c.setFont("Helvetica", 10)
    c.drawString(50, 730, f"Tanggal: {datetime.now().strftime('%d %B %Y %H:%M:%S')}")
    
    c.line(50, 720, 550, 720)
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, 690, "HASIL PERHITUNGAN PSA")
    
    y_position = 670
    c.setFont("Helvetica", 10)
    
    hasil_items = [
        ("Diameter Rata-rata", f"{hasil_psa['diameter_rata']:.2f} nm"),
        ("PDI Rata-rata", f"{hasil_psa['pdi_rata']:.3f}"),
        ("Kualitas Nanomaterial", hasil_psa['kualitas']),
        ("Total Volume", f"{hasil_psa['total_vol']:.1f}%"),
        ("Jumlah Partikel", str(hasil_psa['jumlah_partikel']))
    ]
    
    for item, value in hasil_items:
        c.drawString(70, y_position, f"{item}: {value}")
        y_position -= 20
    
    y_position -= 20
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y_position, "DISTRIBUSI UKURAN PARTIKEL")
    y_position -= 20
    
    c.setFont("Helvetica", 10)
    for ukuran, persentase in hasil_psa['distribusi'].items():
        c.drawString(70, y_position, f"{ukuran}: {persentase:.1f}%")
        y_position -= 15
    
    y_position -= 30
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y_position, "DATA INPUT")
    y_position -= 20
    
    data = [["No", "PDI", "%Vol", "Diameter (nm)"]] + data_input
    table = Table(data, colWidths=[50, 100, 100, 100])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    table.wrapOn(c, 400, 200)
    table.drawOn(c, 50, y_position - len(data_input) * 20)
    
    c.setFont("Helvetica-Oblique", 8)
    c.drawString(50, 30, f"Dibuat dengan NaNote v1.0 • {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

def buat_grafik_distribusi(distribusi):
    fig, ax = plt.subplots(figsize=(8, 5))
    
    labels = list(distribusi.keys())
    values = list(distribusi.values())
    
    bars = ax.bar(labels, values, color=['#3B82F6', '#10B981', '#F59E0B', '#EF4444', '#8B5CF6'])
    
    ax.set_xlabel('Rentang Ukuran', fontsize=12)
    ax.set_ylabel('Persentase Volume (%)', fontsize=12)
    ax.set_title('Distribusi Ukuran Partikel', fontsize=14, fontweight='bold')
    ax.grid(axis='y', alpha=0.3)
    
    for bar in bars:
        height = bar.get_height()
        if height > 0:
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{height:.1f}%', ha='center', va='bottom')
    
    plt.xticks(rotation=45)
    plt.tight_layout()
    return fig

# ============================================
# HALAMAN LOGIN DAN REGISTRASI
# ============================================

def show_login_page(conn):
    """Menampilkan halaman login"""
    st.markdown('<div class="login-container">', unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["Login", "Register"])
    
    with tab1:
        st.markdown('<h3 style="text-align: center;">🔐 Login ke NaNote</h3>', unsafe_allow_html=True)
        
        with st.form("login_form"):
            username = st.text_input("Username atau Email", placeholder="Masukkan username atau email")
            password = st.text_input("Password", type="password", placeholder="Masukkan password")
            
            col1, col2 = st.columns([1, 2])
            with col1:
                remember_me = st.checkbox("Ingat saya")
            
            submitted = st.form_submit_button("Login", use_container_width=True)
            
            if submitted:
                if not username or not password:
                    st.error("Username dan password harus diisi")
                else:
                    success, result = login_user(conn, username, password)
                    
                    if success:
                        st.session_state.user = result
                        st.success("Login berhasil! Mengalihkan...")
                        st.rerun()
                    else:
                        st.error(result)
        
        st.markdown("---")
        st.markdown("""
        **Akun demo:** (untuk testing)
        - Username: `demo`
        - Password: `Demo123!`
        
        **Admin:**
        - Username: `admin`
        - Password: `Admin123!`
        """)
    
    with tab2:
        st.markdown('<h3 style="text-align: center;">📝 Buat Akun Baru</h3>', unsafe_allow_html=True)
        
        with st.form("register_form"):
            col1, col2 = st.columns(2)
            with col1:
                username = st.text_input("Username*", placeholder="Minimal 3 karakter")
                full_name = st.text_input("Nama Lengkap", placeholder="Nama lengkap Anda")
            with col2:
                email = st.text_input("Email*", placeholder="email@contoh.com")
                institution = st.text_input("Institusi", placeholder="Universitas/Perusahaan")
            
            password = st.text_input("Password*", type="password", placeholder="Minimal 8 karakter")
            confirm_password = st.text_input("Konfirmasi Password*", type="password", placeholder="Ulangi password")
            
            # Validasi password strength
            if password:
                is_strong, errors = is_valid_password(password)
                if not is_strong:
                    for error in errors:
                        st.warning(f"⚠️ {error}")
            
            terms = st.checkbox("Saya menyetujui syarat dan ketentuan")
            
            submitted = st.form_submit_button("Daftar", use_container_width=True)
            
            if submitted:
                if not all([username, email, password, confirm_password]):
                    st.error("Semua field bertanda * harus diisi")
                elif password != confirm_password:
                    st.error("Password dan konfirmasi password tidak cocok")
                elif not terms:
                    st.error("Anda harus menyetujui syarat dan ketentuan")
                else:
                    success, message = register_user(conn, username, email, password, full_name, institution)
                    if success:
                        st.success(message)
                        st.info("Silakan login dengan akun yang baru dibuat")
                    else:
                        st.error(message)
        
        st.markdown("---")
        st.markdown("""
        **Persyaratan Password:**
        - Minimal 8 karakter
        - Minimal 1 huruf besar
        - Minimal 1 huruf kecil
        - Minimal 1 angka
        - Minimal 1 karakter khusus
        """)
    
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================
# HALAMAN PROFILE
# ============================================

def show_profile_page(conn, current_user):
    """Menampilkan halaman profile pengguna"""
    st.markdown('<h2 class="section-header">👤 Profil Pengguna</h2>', unsafe_allow_html=True)
    
    # Ambil data lengkap pengguna
    c = conn.cursor()
    c.execute('''
        SELECT u.username, u.email, u.full_name, u.institution, u.role, 
               u.created_at, u.last_login, u.login_count,
               us.page_views, us.catatan_created, us.psa_calculated
        FROM users u
        LEFT JOIN usage_stats us ON u.id = us.user_id
        WHERE u.id = ?
    ''', (current_user['user_id'],))
    
    user_data = c.fetchone()
    
    if user_data:
        username, email, full_name, institution, role, created_at, last_login, login_count, page_views, catatan_created, psa_calculated = user_data
        
        # Tampilkan avatar dan info
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown('<div class="user-profile">', unsafe_allow_html=True)
            st.markdown(f'<div class="user-avatar">{username[0].upper()}</div>', unsafe_allow_html=True)
            st.markdown(f"**{username}**")
            st.markdown(f"*{role}*")
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Tombol logout
            if st.button("🚪 Logout", use_container_width=True):
                logout_user(conn, current_user['user_id'], current_user['session_token'])
                if 'user' in st.session_state:
                    del st.session_state.user
                st.success("Logout berhasil! Mengalihkan...")
                st.rerun()
        
        with col2:
            # Info pengguna
            st.markdown("### Informasi Akun")
            
            info_data = {
                "Username": username,
                "Email": email,
                "Nama Lengkap": full_name or "-",
                "Institusi": institution or "-",
                "Role": role,
                "Akun Dibuat": created_at,
                "Terakhir Login": last_login or "-",
                "Total Login": login_count
            }
            
            for key, value in info_data.items():
                st.markdown(f"**{key}:** {value}")
        
        st.markdown("---")
        
        # Statistik penggunaan
        st.markdown("### 📊 Statistik Penggunaan")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f"""
            <div class="stats-card">
                <h4>📝 Catatan</h4>
                <h2>{catatan_created or 0}</h2>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div class="stats-card">
                <h4>🧮 PSA</h4>
                <h2>{psa_calculated or 0}</h2>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown(f"""
            <div class="stats-card">
                <h4>👁️ Halaman</h4>
                <h2>{page_views or 0}</h2>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown(f"""
            <div class="stats-card">
                <h4>🔑 Login</h4>
                <h2>{login_count or 0}</h2>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Ubah password
        st.markdown("### 🔒 Ubah Password")
        
        with st.expander("Klik untuk mengubah password"):
            with st.form("change_password_form"):
                current_password = st.text_input("Password Saat Ini", type="password")
                new_password = st.text_input("Password Baru", type="password")
                confirm_password = st.text_input("Konfirmasi Password Baru", type="password")
                
                submitted = st.form_submit_button("Ubah Password", use_container_width=True)
                
                if submitted:
                    # Verifikasi password saat ini
                    c.execute("SELECT password_hash, salt FROM users WHERE id = ?", (current_user['user_id'],))
                    db_hash, salt = c.fetchone()
                    
                    if not verify_password(current_password, db_hash, salt):
                        st.error("Password saat ini salah")
                    elif new_password != confirm_password:
                        st.error("Password baru tidak cocok")
                    else:
                        is_strong, errors = is_valid_password(new_password)
                        if not is_strong:
                            st.error(f"Password lemah: {', '.join(errors)}")
                        else:
                            # Hash password baru
                            new_hash, new_salt = hash_password(new_password)
                            
                            c.execute('''
                                UPDATE users 
                                SET password_hash = ?, salt = ?
                                WHERE id = ?
                            ''', (new_hash, new_salt, current_user['user_id']))
                            
                            conn.commit()
                            log_activity(conn, current_user['user_id'], 'password_changed', {})
                            st.success("Password berhasil diubah!")

# ============================================
# HALAMAN ADMIN
# ============================================

def show_admin_page(conn, current_user):
    """Halaman admin untuk mengelola pengguna"""
    if current_user['role'] != 'admin':
        st.error("Akses ditolak. Hanya admin yang dapat mengakses halaman ini.")
        return
    
    st.markdown('<h2 class="section-header">👑 Admin Dashboard</h2>', unsafe_allow_html=True)
    
    tab1, tab2, tab3 = st.tabs(["Manajemen Pengguna", "Statistik Sistem", "Log Aktivitas"])
    
    with tab1:
        st.markdown("### 📋 Daftar Pengguna")
        
        # Ambil semua pengguna
        c = conn.cursor()
        c.execute('''
            SELECT u.id, u.username, u.email, u.full_name, u.role, u.is_active, 
                   u.created_at, u.last_login, u.login_count,
                   us.page_views, us.catatan_created, us.psa_calculated
            FROM users u
            LEFT JOIN usage_stats us ON u.id = us.user_id
            ORDER BY u.created_at DESC
        ''')
        
        users = c.fetchall()
        
        if users:
            # Tampilkan dalam dataframe
            df_users = pd.DataFrame(users, columns=[
                'ID', 'Username', 'Email', 'Nama', 'Role', 'Aktif', 
                'Dibuat', 'Login Terakhir', 'Jumlah Login',
                'Page Views', 'Catatan', 'PSA'
            ])
            
            # Konversi tipe data
            df_users['Aktif'] = df_users['Aktif'].map({1: '✅', 0: '❌'})
            
            st.dataframe(df_users, use_container_width=True)
            
            # Fitur edit pengguna
            st.markdown("### ✏️ Edit Pengguna")
            col1, col2 = st.columns(2)
            
            with col1:
                user_ids = [str(user[0]) for user in users]
                selected_user = st.selectbox("Pilih Pengguna", user_ids)
                
                if selected_user:
                    c.execute("SELECT role, is_active FROM users WHERE id = ?", (selected_user,))
                    role, is_active = c.fetchone()
                    
                    new_role = st.selectbox(
                        "Role",
                        ["user", "admin"],
                        index=0 if role == "user" else 1
                    )
                    
                    new_status = st.selectbox(
                        "Status",
                        ["Aktif", "Nonaktif"],
                        index=0 if is_active == 1 else 1
                    )
                    
                    if st.button("💾 Update Pengguna", use_container_width=True):
                        c.execute('''
                            UPDATE users 
                            SET role = ?, is_active = ?
                            WHERE id = ?
                        ''', (new_role, 1 if new_status == "Aktif" else 0, selected_user))
                        conn.commit()
                        st.success("Pengguna berhasil diupdate!")
                        st.rerun()
            
            with col2:
                st.markdown("### 🆕 Tambah Pengguna")
                with st.form("add_user_form"):
                    new_username = st.text_input("Username")
                    new_email = st.text_input("Email")
                    new_password = st.text_input("Password", type="password")
                    new_role = st.selectbox("Role", ["user", "admin"])
                    
                    submitted = st.form_submit_button("Tambah Pengguna", use_container_width=True)
                    
                    if submitted:
                        if not all([new_username, new_email, new_password]):
                            st.error("Semua field harus diisi")
                        else:
                            success, message = register_user(conn, new_username, new_email, new_password, role=new_role)
                            if success:
                                st.success(message)
                                st.rerun()
                            else:
                                st.error(message)
        
        else:
            st.info("Belum ada pengguna terdaftar")
    
    with tab2:
        st.markdown("### 📈 Statistik Sistem")
        
        c = conn.cursor()
        
        # Statistik umum
        c.execute("SELECT COUNT(*) FROM users")
        total_users = c.fetchone()[0]
        
        c.execute("SELECT COUNT(*) FROM users WHERE is_active = 1")
        active_users = c.fetchone()[0]
        
        c.execute("SELECT COUNT(*) FROM catatan_praktikum")
        total_catatan = c.fetchone()[0]
        
        c.execute("SELECT COUNT(*) FROM hasil_psa")
        total_psa = c.fetchone()[0]
        
        c.execute("SELECT COUNT(*) FROM user_history")
        total_activities = c.fetchone()[0]
        
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.metric("Total Pengguna", total_users)
        with col2:
            st.metric("Pengguna Aktif", active_users)
        with col3:
            st.metric("Total Catatan", total_catatan)
        with col4:
            st.metric("Total PSA", total_psa)
        with col5:
            st.metric("Aktivitas", total_activities)
        
        # Grafik aktivitas
        st.markdown("### 📊 Aktivitas Harian (7 Hari Terakhir)")
        
        c.execute('''
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM user_history
            WHERE timestamp >= datetime('now', '-7 days')
            GROUP BY DATE(timestamp)
            ORDER BY date
        ''')
        
        activity_data = c.fetchall()
        
        if activity_data:
            dates = [row[0] for row in activity_data]
            counts = [row[1] for row in activity_data]
            
            fig, ax = plt.subplots(figsize=(10, 4))
            ax.bar(dates, counts, color='#3B82F6')
            ax.set_xlabel('Tanggal')
            ax.set_ylabel('Jumlah Aktivitas')
            ax.set_title('Aktivitas 7 Hari Terakhir')
            plt.xticks(rotation=45)
            plt.tight_layout()
            st.pyplot(fig)
    
    with tab3:
        st.markdown("### 📋 Log Aktivitas Sistem")
        
        c = conn.cursor()
        c.execute('''
            SELECT uh.timestamp, u.username, uh.activity_type, uh.activity_data, uh.ip_address
            FROM user_history uh
            JOIN users u ON uh.user_id = u.id
            ORDER BY uh.timestamp DESC
            LIMIT 100
        ''')
        
        logs = c.fetchall()
        
        if logs:
            df_logs = pd.DataFrame(logs, columns=['Waktu', 'Pengguna', 'Tipe Aktivitas', 'Data', 'IP Address'])
            
            # Filter berdasarkan tipe aktivitas
            activity_types = df_logs['Tipe Aktivitas'].unique()
            selected_types = st.multiselect("Filter Tipe Aktivitas", activity_types, default=activity_types[:5])
            
            if selected_types:
                df_filtered = df_logs[df_logs['Tipe Aktivitas'].isin(selected_types)]
                st.dataframe(df_filtered, use_container_width=True)
                
                # Download log
                csv = df_filtered.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📥 Download Log (CSV)",
                    data=csv,
                    file_name=f"system_logs_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            else:
                st.info("Pilih tipe aktivitas untuk menampilkan log")
        else:
            st.info("Belum ada log aktivitas")

# ============================================
# FUNGSI DATABASE UNTUK APLIKASI UTAMA
# ============================================

def save_catatan_to_db(conn, user_id, catatan_data):
    """Menyimpan catatan ke database"""
    c = conn.cursor()
    
    try:
        c.execute('''
            INSERT INTO catatan_praktikum 
            (user_id, judul, praktikan, mata_praktikum, tanggal, kelompok, pic, 
             tujuan, alat_bahan, prosedur, hasil, analisis, kesimpulan)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            user_id,
            catatan_data['judul'],
            catatan_data['praktikan'],
            catatan_data['mata_praktikum'],
            catatan_data['tanggal'],
            catatan_data.get('kelompok', ''),
            catatan_data.get('pic', ''),
            catatan_data['isi']['tujuan'],
            catatan_data['isi']['alat_bahan'],
            catatan_data['isi']['prosedur'],
            catatan_data['isi']['hasil'],
            catatan_data['isi'].get('analisis', ''),
            catatan_data['isi'].get('kesimpulan', '')
        ))
        
        catatan_id = c.lastrowid
        
        # Update statistik
        c.execute('''
            UPDATE usage_stats 
            SET catatan_created = catatan_created + 1 
            WHERE user_id = ?
        ''', (user_id,))
        
        conn.commit()
        
        # Log aktivitas
        log_activity(conn, user_id, 'catatan_created', {
            'catatan_id': catatan_id,
            'judul': catatan_data['judul']
        })
        
        return catatan_id
    
    except Exception as e:
        return None

def save_psa_to_db(conn, user_id, psa_data, data_input):
    """Menyimpan hasil PSA ke database"""
    c = conn.cursor()
    
    try:
        c.execute('''
            INSERT INTO hasil_psa 
            (user_id, data_input, diameter_rata, pdi_rata, total_vol, 
             kualitas, distribusi, jumlah_partikel)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            user_id,
            json.dumps(data_input),
            psa_data['diameter_rata'],
            psa_data['pdi_rata'],
            psa_data['total_vol'],
            psa_data['kualitas'],
            json.dumps(psa_data['distribusi']),
            psa_data['jumlah_partikel']
        ))
        
        psa_id = c.lastrowid
        
        # Update statistik
        c.execute('''
            UPDATE usage_stats 
            SET psa_calculated = psa_calculated + 1 
            WHERE user_id = ?
        ''', (user_id,))
        
        conn.commit()
        
        # Log aktivitas
        log_activity(conn, user_id, 'psa_calculated', {
            'psa_id': psa_id,
            'diameter_rata': psa_data['diameter_rata']
        })
        
        return psa_id
    
    except Exception as e:
        return None

def get_user_catatan(conn, user_id):
    """Mendapatkan catatan pengguna dari database"""
    c = conn.cursor()
    
    c.execute('''
        SELECT * FROM catatan_praktikum 
        WHERE user_id = ? 
        ORDER BY created_at DESC
    ''', (user_id,))
    
    rows = c.fetchall()
    catatan_list = []
    
    for row in rows:
        catatan_list.append({
            'id': row[0],
            'user_id': row[1],
            'judul': row[2],
            'praktikan': row[3],
            'mata_praktikum': row[4],
            'tanggal': row[5],
            'kelompok': row[6],
            'pic': row[7],
            'isi': {
                'tujuan': row[8],
                'alat_bahan': row[9],
                'prosedur': row[10],
                'hasil': row[11],
                'analisis': row[12],
                'kesimpulan': row[13]
            },
            'is_public': row[14],
            'created_at': row[15],
            'updated_at': row[16]
        })
    
    return catatan_list

def get_user_psa_history(conn, user_id):
    """Mendapatkan riwayat PSA pengguna dari database"""
    c = conn.cursor()
    
    c.execute('''
        SELECT * FROM hasil_psa 
        WHERE user_id = ? 
        ORDER BY created_at DESC
    ''', (user_id,))
    
    rows = c.fetchall()
    psa_list = []
    
    for row in rows:
        psa_list.append({
            'id': row[0],
            'user_id': row[1],
            'data_input': json.loads(row[2]),
            'diameter_rata': row[3],
            'pdi_rata': row[4],
            'total_vol': row[5],
            'kualitas': row[6],
            'distribusi': json.loads(row[7]),
            'jumlah_partikel': row[8],
            'is_public': row[9],
            'created_at': row[10]
        })
    
    return psa_list

# ============================================
# APLIKASI UTAMA
# ============================================

def main_app(conn, current_user):
    """Aplikasi utama setelah login"""
    
    # Inisialisasi session state
    if 'catatan_list' not in st.session_state:
        st.session_state.catatan_list = []
    if 'psa_data' not in st.session_state:
        st.session_state.psa_data = []
    if 'current_note' not in st.session_state:
        st.session_state.current_note = {}
    if 'current_psa' not in st.session_state:
        st.session_state.current_psa = {}
    if 'show_download' not in st.session_state:
        st.session_state.show_download = False
    
    # Sidebar dengan informasi pengguna
    with st.sidebar:
        st.image("https://i.pinimg.com/1200x/8b/06/a8/8b06a832394c6d214729546d6888d0d0.jpg", width=80)
        st.title("NaNote")
        st.markdown("**Catatan & Kalkulator PSA**")
        
        # Info pengguna
        st.markdown("---")
        st.markdown(f"👋 **Halo, {current_user['username']}!**")
        if current_user.get('full_name'):
            st.markdown(f"*{current_user['full_name']}*")
        
        # Menu navigasi
        st.markdown("---")
        
        menu_items = ["🏠 Beranda", "📝 Catatan Praktikum", "🧮 Kalkulasi PSA", "📊 Data Tersimpan", "👤 Profil"]
        
        # Tambahkan menu admin jika role admin
        if current_user['role'] == 'admin':
            menu_items.append("👑 Admin")
        
        menu_items.append("ℹ️ Panduan")
        
        menu = st.radio(
            "Pilih Menu:",
            menu_items
        )
        
        # Log aktivitas page view
        log_activity(conn, current_user['user_id'], 'page_view', {
            'page': menu,
            'username': current_user['username']
        })
    
    # Konten utama berdasarkan menu
    if menu == "🏠 Beranda":
        st.markdown('<h1 class="main-title">🔬 NaNote</h1>', unsafe_allow_html=True)
        st.markdown('<p class="sub-title">Aplikasi Catatan Praktikum & Kalkulator PSA untuk Nanomaterial</p>', unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            <div class="data-box">
                <h3>📝 Catatan Praktikum</h3>
                <p>Buat dan simpan catatan praktikum Anda dalam format Microsoft Word (.docx).</p>
                <ul>
                    <li>Editor teks lengkap</li>
                    <li>Template otomatis</li>
                    <li>Simpan sebagai .docx</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("""
            <div class="data-box">
                <h3>🧮 Kalkulasi PSA</h3>
                <p>Hitung hasil Particle Size Analysis dari data PDI, %vol, dan diameter.</p>
                <ul>
                    <li>Input data multiple</li>
                    <li>Perhitungan otomatis</li>
                    <li>Analisis kualitas</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown("""
            <div class="data-box">
                <h3>📊 Ekspor Data</h3>
                <p>Ekspor hasil dalam format standar untuk laporan.</p>
                <ul>
                    <li>Catatan → Word (.docx)</li>
                    <li>Hasil PSA → PDF</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Statistik pengguna
        st.markdown("### 📈 Statistik Anda")
        
        c = conn.cursor()
        c.execute('''
            SELECT catatan_created, psa_calculated, page_views 
            FROM usage_stats 
            WHERE user_id = ?
        ''', (current_user['user_id'],))
        
        stats = c.fetchone()
        
        if stats:
            catatan_created, psa_calculated, page_views = stats
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown(f"""
                <div class="stats-card">
                    <h4>📝 Catatan</h4>
                    <h2>{catatan_created or 0}</h2>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="stats-card">
                    <h4>🧮 PSA</h4>
                    <h2>{psa_calculated or 0}</h2>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown(f"""
                <div class="stats-card">
                    <h4>👁️ Halaman</h4>
                    <h2>{page_views or 0}</h2>
                </div>
                """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown("### 🎯 Fitur Utama")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            **Untuk Catatan Praktikum:**
            - Input data praktikum lengkap
            - Kategorisasi otomatis
            - Template siap pakai
            - Ekspor ke Microsoft Word
            - Penyimpanan database
            """)
        
        with col2:
            st.markdown("""
            **Untuk Kalkulasi PSA:**
            - Input data PDI, %vol, diameter
            - Perhitungan rata-rata berbobot
            - Analisis distribusi ukuran
            - Penilaian kualitas nanomaterial
            - Ekspor ke PDF profesional
            """)
    
    elif menu == "📝 Catatan Praktikum":
        st.markdown('<h2 class="section-header">📝 Buat Catatan Praktikum</h2>', unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["Buat Catatan Baru", "Lihat Catatan Tersimpan"])
        
        with tab1:
            with st.form("catatan_form"):
                col1, col2 = st.columns(2)
                
                with col1:
                    judul = st.text_input("Judul Catatan*", placeholder="Contoh: Praktikum Sintesis Nanopartikel Fe3O4")
                    praktikan = st.text_input("Nama Praktikan*", placeholder="Nama lengkap Anda")
                    mata_praktikum = st.text_input("Mata Praktikum*", placeholder="Contoh: Nanomaterial 2")
                
                with col2:
                    tanggal = st.date_input("Tanggal Praktikum*", datetime.now())
                    kelompok = st.text_input("Kelompok", placeholder="Contoh: Kelompok 5")
                    pic = st.text_input("PIC Praktikum", placeholder="Nama PIC Praktikum")
                
                st.markdown("### Isi Catatan")
                
                col1, col2 = st.columns(2)
                with col1:
                    tujuan = st.text_area("Tujuan Praktikum*", height=150, 
                                        placeholder="Tuliskan tujuan praktikum...")
                    alat_bahan = st.text_area("Alat dan Bahan*", height=150,
                                            placeholder="Daftar alat dan bahan...")
                
                with col2:
                    prosedur = st.text_area("Prosedur Kerja*", height=150,
                                          placeholder="Langkah-langkah percobaan...")
                    hasil = st.text_area("Hasil Pengamatan*", height=150,
                                       placeholder="Hasil yang diamati...")
                
                analisis = st.text_area("Analisis Data", height=120,
                                      placeholder="Analisis dari hasil yang diperoleh...")
                kesimpulan = st.text_area("Kesimpulan", height=100,
                                        placeholder="Kesimpulan praktikum...")
                
                submitted = st.form_submit_button("💾 Simpan Catatan", use_container_width=True)
                
                if submitted:
                    if not all([judul, praktikan, mata_praktikum, tujuan, 
                              alat_bahan, prosedur, hasil]):
                        st.error("Harap isi semua field yang wajib diisi (*)")
                    else:
                        catatan_data = {
                            'judul': judul,
                            'praktikan': praktikan,
                            'mata_praktikum': mata_praktikum,
                            'tanggal': tanggal.strftime("%Y-%m-%d"),
                            'kelompok': kelompok,
                            'pic': pic,
                            'isi': {
                                'tujuan': tujuan,
                                'alat_bahan': alat_bahan,
                                'prosedur': prosedur,
                                'hasil': hasil,
                                'analisis': analisis,
                                'kesimpulan': kesimpulan
                            }
                        }
                        
                        # Simpan ke database
                        catatan_id = save_catatan_to_db(conn, current_user['user_id'], catatan_data)
                        
                        if catatan_id:
                            st.success(f"✅ Catatan berhasil disimpan! ID: {catatan_id}")
                            
                            # Update session state
                            catatan_data['id'] = catatan_id
                            catatan_data['waktu_buat'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            st.session_state.current_note = catatan_data
                            st.session_state.show_download = True
                        else:
                            st.error("❌ Gagal menyimpan catatan ke database")
                        
                        # Tampilkan preview
                        with st.expander("Preview Catatan", expanded=True):
                            st.markdown(f"**Judul:** {judul}")
                            st.markdown(f"**Praktikan:** {praktikan} | **Tanggal:** {tanggal}")
                            st.markdown("---")
                            st.markdown(f"**Tujuan:**\n{tujuan}")
                            st.markdown(f"**Alat dan Bahan:**\n{alat_bahan}")
                            st.markdown(f"**Prosedur:**\n{prosedur}")
                            st.markdown(f"**Hasil:**\n{hasil}")
            
            # Tombol download
            if st.session_state.show_download and st.session_state.current_note:
                catatan_data = st.session_state.current_note
                doc = buat_file_word(catatan_data)
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                st.download_button(
                    label="⬇️ Download sebagai Word (.docx)",
                    data=doc_buffer,
                    file_name=f"catatan_{catatan_data['judul'].replace(' ', '_')}_{catatan_data['tanggal']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with tab2:
            # Ambil data dari database
            catatan_list = get_user_catatan(conn, current_user['user_id'])
            
            if not catatan_list:
                st.info("📝 Belum ada catatan yang disimpan.")
            else:
                st.markdown(f"### 📚 Catatan Tersimpan ({len(catatan_list)})")
                
                for catatan in catatan_list:
                    with st.expander(f"{catatan['judul']} - {catatan['tanggal']} (ID: {catatan['id']})"):
                        col1, col2, col3 = st.columns([3, 1, 1])
                        
                        with col1:
                            st.markdown(f"**Praktikan:** {catatan['praktikan']}")
                            st.markdown(f"**Mata Praktikum:** {catatan['mata_praktikum']}")
                            st.markdown(f"**Dibuat:** {catatan['created_at']}")
                            if catatan['kelompok']:
                                st.markdown(f"**Kelompok:** {catatan['kelompok']}")
                        
                        with col2:
                            # Tombol download
                            doc = buat_file_word(catatan)
                            doc_buffer = io.BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Download",
                                data=doc_buffer,
                                file_name=f"catatan_{catatan['id']}_{catatan['judul'].replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_{catatan['id']}"
                            )
                        
                        with col3:
                            if st.button("🗑️ Hapus", key=f"delete_{catatan['id']}"):
                                c = conn.cursor()
                                c.execute("DELETE FROM catatan_praktikum WHERE id = ? AND user_id = ?", 
                                         (catatan['id'], current_user['user_id']))
                                conn.commit()
                                st.success("✅ Catatan berhasil dihapus!")
                                st.rerun()
    
    elif menu == "🧮 Kalkulasi PSA":
        st.markdown('<h2 class="section-header">🧮 Kalkulasi Particle Size Analysis</h2>', unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["Input Data", "Hasil Kalkulasi"])
        
        with tab1:
            st.markdown("""
            <div class="warning-box">
            <strong>📋 Panduan Input Data:</strong><br>
            1. PDI (Polydispersity Index): 0.0 - 1.0 (semakin kecil semakin baik)<br>
            2. %Vol: Persentase volume partikel (0-100%)<br>
            3. Diameter: Ukuran partikel dalam nanometer (nm)<br>
            4. Minimal 2 data untuk perhitungan
            </div>
            """, unsafe_allow_html=True)
            
            # Input jumlah data
            col1, col2 = st.columns(2)
            with col1:
                jumlah_data = st.number_input(
                    "Jumlah Data:",
                    min_value=2,
                    max_value=20,
                    value=5,
                    step=1
                )
            
            with col2:
                st.markdown("")
                st.markdown("")
                if st.button("🔄 Reset Tabel", use_container_width=True):
                    st.session_state.psa_data = []
                    st.rerun()
            
            # Input data dalam bentuk tabel
            st.markdown("### Masukkan Data PSA")
            
            data_input = []
            for i in range(jumlah_data):
                cols = st.columns(4)
                with cols[0]:
                    ulangan = st.text_input(f"Ulangan", value=str(i+1), disabled=True, key=f"ulangan_{i}")
                with cols[1]:
                    pdi = st.number_input(
                        f"PDI",
                        min_value=0.0,
                        max_value=1.0,
                        value=0.15,
                        step=0.01,
                        key=f"pdi_{i}"
                    )
                with cols[2]:
                    vol = st.number_input(
                        f"%Vol",
                        min_value=0.0,
                        max_value=100.0,
                        value=20.0,
                        step=1.0,
                        key=f"vol_{i}"
                    )
                with cols[3]:
                    diameter = st.number_input(
                        f"Diameter (nm)",
                        min_value=0.1,
                        value=50.0,
                        step=1.0,
                        key=f"diameter_{i}"
                    )
                data_input.append([pdi, vol, diameter])
            
            # Tombol hitung
            if st.button("🚀 Kalkulasi Hasil PSA", type="primary", use_container_width=True):
                # Ekstrak data
                pdi_list = [d[0] for d in data_input]
                vol_list = [d[1] for d in data_input]
                diameter_list = [d[2] for d in data_input]
                
                # Kalkulasi hasil
                hasil = kalkulasi_hasil_psa(pdi_list, vol_list, diameter_list)
                
                if hasil:
                    # Simpan ke database
                    psa_id = save_psa_to_db(conn, current_user['user_id'], hasil, data_input)
                    
                    if psa_id:
                        st.success(f"✅ Hasil PSA berhasil disimpan! ID: {psa_id}")
                        
                        st.session_state.current_psa = {
                            'id': psa_id,
                            'hasil': hasil,
                            'data_input': data_input,
                            'pdi_list': pdi_list,
                            'vol_list': vol_list,
                            'diameter_list': diameter_list
                        }
                    else:
                        st.error("❌ Gagal menyimpan hasil PSA ke database")
                        st.session_state.current_psa = {
                            'hasil': hasil,
                            'data_input': data_input,
                            'pdi_list': pdi_list,
                            'vol_list': vol_list,
                            'diameter_list': diameter_list
                        }
                    
                    st.success("✅ Perhitungan selesai! Buka tab 'Hasil Kalkulasi'.")
                    st.rerun()
        
        with tab2:
            if 'current_psa' not in st.session_state or not st.session_state.current_psa:
                st.info("ℹ️ Masukkan data di tab 'Input Data' dan klik 'Kalkulasi Hasil PSA'.")
            else:
                hasil = st.session_state.current_psa['hasil']
                data_input = st.session_state.current_psa['data_input']
                
                st.markdown("### 📊 Hasil Kalkulasi")
                
                # Tampilkan ID jika ada
                if 'id' in st.session_state.current_psa:
                    st.caption(f"ID: {st.session_state.current_psa['id']}")
                
                # Tampilkan hasil utama
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown(f"""
                    <div class="result-box">
                    <h4>📐 Diameter Rata-rata</h4>
                    <h3>{hasil['diameter_rata']:.2f} nm</h3>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    st.markdown(f"""
                    <div class="result-box">
                    <h4>📊 PDI Rata-rata</h4>
                    <h3>{hasil['pdi_rata']:.3f}</h3>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col3:
                    warna = hasil['warna_kualitas']
                    st.markdown(f"""
                    <div class="result-box">
                    <h4>🏆 Kualitas</h4>
                    <h3 style='color: {warna};'>{hasil['kualitas']}</h3>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Tabel data input
                st.markdown("### 📋 Data Input")
                df_input = pd.DataFrame({
                    'No': range(1, len(data_input) + 1),
                    'PDI': [f"{d[0]:.3f}" for d in data_input],
                    '%Vol': [f"{d[1]:.1f}%" for d in data_input],
                    'Diameter (nm)': [f"{d[2]:.1f}" for d in data_input]
                })
                st.dataframe(df_input, use_container_width=True)
                
                # Grafik distribusi
                st.markdown("### 📈 Distribusi Ukuran Partikel")
                fig = buat_grafik_distribusi(hasil['distribusi'])
                st.pyplot(fig)
                
                # Tabel distribusi
                st.markdown("### 📊 Detail Distribusi")
                df_distribusi = pd.DataFrame({
                    'Rentang Ukuran': list(hasil['distribusi'].keys()),
                    'Persentase Volume': [f"{v:.1f}%" for v in hasil['distribusi'].values()]
                })
                st.dataframe(df_distribusi, use_container_width=True)
                
                # Tombol download PDF
                pdf_buffer = buat_file_pdf(hasil, df_input.values.tolist())
                
                st.download_button(
                    label="📄 Download Hasil sebagai PDF",
                    data=pdf_buffer,
                    file_name=f"hasil_PSA_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key="download_pdf"
                )
    
    elif menu == "📊 Data Tersimpan":
        st.markdown('<h2 class="section-header">📊 Data Tersimpan</h2>', unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["Catatan Praktikum", "Hasil PSA"])
        
        with tab1:
            catatan_list = get_user_catatan(conn, current_user['user_id'])
            
            if not catatan_list:
                st.info("📝 Belum ada catatan yang disimpan.")
            else:
                st.markdown(f"### Total Catatan: {len(catatan_list)}")
                
                for catatan in catatan_list:
                    with st.expander(f"{catatan['judul']} ({catatan['tanggal']}) - ID: {catatan['id']}"):
                        col1, col2, col3 = st.columns([3, 1, 1])
                        
                        with col1:
                            st.markdown(f"**Praktikan:** {catatan['praktikan']}")
                            st.markdown(f"**Mata Praktikum:** {catatan['mata_praktikum']}")
                            st.markdown(f"**Dibuat:** {catatan['created_at']}")
                        
                        with col2:
                            # Tombol download
                            doc = buat_file_word(catatan)
                            doc_buffer = io.BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Word",
                                data=doc_buffer,
                                file_name=f"catatan_{catatan['id']}_{catatan['judul'].replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"word_download_{catatan['id']}"
                            )
                        
                        with col3:
                            if st.button("🗑️ Hapus", key=f"del_cat_{catatan['id']}"):
                                c = conn.cursor()
                                c.execute("DELETE FROM catatan_praktikum WHERE id = ? AND user_id = ?", 
                                         (catatan['id'], current_user['user_id']))
                                conn.commit()
                                st.success("✅ Catatan berhasil dihapus!")
                                st.rerun()
                
                # Tombol backup
                st.markdown("---")
                if st.button("💾 Backup Semua Catatan", use_container_width=True):
                    backup_data = {
                        'catatan': catatan_list,
                        'backup_time': datetime.now().isoformat(),
                        'total_catatan': len(catatan_list)
                    }
                    
                    backup_json = json.dumps(backup_data, indent=2, ensure_ascii=False)
                    
                    st.download_button(
                        label="📥 Download Backup JSON",
                        data=backup_json,
                        file_name=f"nanote_backup_catatan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        use_container_width=True
                    )
        
        with tab2:
            psa_list = get_user_psa_history(conn, current_user['user_id'])
            
            if not psa_list:
                st.info("🧮 Belum ada hasil PSA yang disimpan.")
            else:
                st.markdown(f"### Total Hasil PSA: {len(psa_list)}")
                
                for psa in psa_list:
                    with st.expander(f"PSA {psa['id']} - {psa['created_at']}"):
                        col1, col2, col3 = st.columns([3, 1, 1])
                        
                        with col1:
                            st.markdown(f"**Waktu:** {psa['created_at']}")
                            st.markdown(f"**Diameter Rata-rata:** {psa['diameter_rata']:.2f} nm")
                            st.markdown(f"**PDI Rata-rata:** {psa['pdi_rata']:.3f}")
                            st.markdown(f"**Kualitas:** {psa['kualitas']}")
                            st.markdown(f"**Jumlah Partikel:** {psa['jumlah_partikel']}")
                        
                        with col2:
                            if st.button("👁️ Detail", key=f"view_psa_{psa['id']}"):
                                st.session_state.current_psa = {
                                    'id': psa['id'],
                                    'hasil': psa,
                                    'data_input': psa['data_input']
                                }
                                st.rerun()
                        
                        with col3:
                            if st.button("🗑️ Hapus", key=f"del_psa_{psa['id']}"):
                                c = conn.cursor()
                                c.execute("DELETE FROM hasil_psa WHERE id = ? AND user_id = ?", 
                                         (psa['id'], current_user['user_id']))
                                conn.commit()
                                st.success("✅ Data PSA berhasil dihapus!")
                                st.rerun()
    
    elif menu == "👤 Profil":
        show_profile_page(conn, current_user)
    
    elif menu == "👑 Admin":
        show_admin_page(conn, current_user)
    
    else:  # Panduan
        st.markdown('<h2 class="section-header">ℹ️ Panduan Penggunaan</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        ### 📚 Tentang NaNote
        
        NaNote adalah aplikasi web yang dirancang khusus untuk membantu praktikan dalam:
        1. **Mencatat hasil praktikum** nanomaterial secara digital
        2. **Menganalisis data PSA** (Particle Size Analysis)
        3. **Menyimpan dan mengekspor** hasil dalam format standar
        
        ### 🔎 Cara Menggunakan
        
        #### 1. Catatan Praktikum
        - Buka menu **"📝 Catatan Praktikum"**
        - Isi form dengan data lengkap praktikum
        - Simpan catatan dan download sebagai file Word (.docx)
        
        #### 2. Kalkulator PSA
        - Buka menu **"🧮 Kalkulator PSA"**
        - Input data PDI, %vol, dan diameter untuk setiap partikel
        - Klik "Hitung Hasil PSA"
        - Lihat hasil di tab "Hasil Perhitungan"
        - Download hasil sebagai PDF
        
        #### 3. Data Tersimpan
        - Buka menu **"📊 Data Tersimpan"**
        - Lihat semua catatan dan hasil perhitungan yang telah disimpan
        - Ekspor data backup jika diperlukan
        
        ### 📊 Interpretasi Hasil PSA
        
        **Polydispersity Index (PDI):**
        - **< 0.1**: Sangat Baik (Monodispers)
        - **0.1 - 0.2**: Baik
        - **0.2 - 0.3**: Cukup
        - **> 0.3**: Kurang (Polydispers Tinggi)
        
        **Diameter:**
        - < 10 nm: Ultra kecil
        - 10-50 nm: Kecil
        - 50-100 nm: Sedang
        - 100-500 nm: Besar
        - > 500 nm: Sangat besar
        
        ### 🔒 Keamanan Akun
        - Gunakan password yang kuat
        - Jangan bagikan akun Anda
        - Logout setelah selesai menggunakan
        - Ubah password secara berkala
        
        ### 🛠️ Teknologi
        - **Framework**: Streamlit (Python)
        - **Database**: SQLite dengan autentikasi
        - **Format Ekspor**: .docx, .pdf, .json
        - **Deployment**: Streamlit Cloud
        - **Bahasa**: Indonesia
        """)
    
    # Footer
    st.markdown("---")
    st.markdown('<div class="footer">', unsafe_allow_html=True)
    st.markdown("🔬 **NaNote** • Aplikasi Catatan Praktikum & Kalkulasi PSA • Dibuat oleh Kelompok 3 Logika dan Pemrograman Komputer")
    st.markdown(f"© {datetime.now().year} • Versi: 2.0 • User: {current_user['username']}")
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================
# MAIN EXECUTION
# ============================================

def main():
    """Fungsi utama aplikasi"""
    # Inisialisasi database
    conn = init_database()
    
    # Cek apakah pengguna sudah login
    current_user = get_current_user(conn)
    
    if current_user:
        # Tampilkan aplikasi utama
        main_app(conn, current_user)
    else:
        # Tampilkan halaman login
        show_login_page(conn)
        
        # Footer untuk halaman login
        st.markdown("---")
        st.markdown('<div class="footer">', unsafe_allow_html=True)
        st.markdown("🔬 **NaNote** • Aplikasi Catatan Praktikum & Kalkulasi PSA")
        st.markdown(f"© {datetime.now().year} • Versi 2.0 dengan Sistem Login")
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Tutup koneksi
    conn.close()

if __name__ == "__main__":
    main()
